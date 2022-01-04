from requests.models import to_native_string
from selenium import webdriver
import requests,bs4,os,re
import docx
import shutil
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import *
from docx.shared import RGBColor
from docx.shared import Pt
from flask import * 
browser= ''
URL = ''
flag = 1
options = webdriver.ChromeOptions()
options.add_experimental_option('excludeSwitches', ['enable-logging'])
options.add_argument("--headless")



def GetSections(essay): # FUNCTION TO GET THE CONTENTS OF THE TOPIC SELECTED

    global URL,browser

    # GETTING TOPIC FROM THE USER FOR THE ESSAY AND USING SELENIUM TO SEARCH WIKIPEDIA
    browser= webdriver.Chrome(options=options)
    browser.get('https://en.wikipedia.org/wiki/Main_Page')
    elem = browser.find_element_by_name('search')
    elem.send_keys(essay)
    elem = browser.find_element_by_id('searchButton')
    elem.click()


    # GETTING THE NEW RESULTANT URL AND PASSING IT FOR WEB SCRAPING
    URL = browser.current_url
    Url = requests.get(URL)
    soup = bs4.BeautifulSoup(Url.text, features="html.parser")



    # USING THE CONTENT INDEX IN WIKIPEDIA TO DISPLAY ALL AVAILABLE TOPICS
    i = 0
    flag  = 0
    sections = []
    max = 0
    elem = soup.select('.toclevel-1') # TO SELECT ONLY THE MAIN WIKIPEDIA SECTIONS
    soup = bs4.BeautifulSoup(str(elem), features="html.parser") 
    while max<10:
            i = i + 1
            j = str(i)
            elem2 = soup.select('.tocsection-'+j+' .tocnumber') # TO GET THE SECTION NUMBER AND TEXT
            if(elem2 == []):
                break
            if(not (elem2[0].text).isdecimal()):
                continue        
            item = elem2[0].text
            max = int(elem2[0].text) # MAINTAINS THE NUMBER OF SECTIONS
            elem2 = soup.select('.tocsection-'+j+' .toctext')
            #DONE TO GET THE URL EQUIVALENT OF THE SECTIONS FOR APPENDING LATER
            data = elem2[0].text
            data = data.split(' ')
            data = ('_').join(data)
            sections = sections + [data]
            item += " "+elem2[0].text
    
    #RETURNS MAX AND SECTION LIST      
    return max, sections

def GetSelection(max,sections,select1): # FUNCTION TO FIND THE SELECTED SECTIONS FROM THE INPUT STRING
    select = [] # LIST TO STORE THE SELECTED INDEXES
    try: # EXCEPTION HANDLING FOR ANY ERRORS
        if ',' in select1:
            choice = select1.split(',') # INITIAL SPLITING
            for i in choice:
                if '-' in i:
                    choice1 = i.split('-')
                    for j in range(int(choice1[0]),int(choice1[1])+1):
                        select += [j]
                else:
                    select  += [int(i)]
        elif '-' in select1: # SPLITING FOR RANGE INPUTS
            choice2 = select1.split('-')
            for i in range(int(choice2[0]),int(choice2[1])+1):
                select += [i]
        else:
            select += [int(select1)]
        for i in select: # CHECKING ERROR CONDITIONS AND CHECK FOR 0
            if i>max or i<0:
                return [-1];
            if i == 0 :
                select = []
                for i in range(1,max+1):
                    select += [i]
        AddSections = [] # CREATING A LIST OF ALL THE SELECTED SECTIONS
        for i in select:
            AddSections += [sections[i-1]]
        
        return AddSections # RETURNING THE LIST OF SECTIONS SELECTED
    except:
        return [-1]; # ERROR MESSAGE FOR EXCEPTION HANDLING

def GetParagraphs(AddSections): # FUNCTION TO GET THE DATA FOR THE SELECTED SECTIONS
    headings = []
    paragraphs = []
    try: # EXCEPTION HANDLING
        for i in AddSections:
            Url = requests.get(URL+'#'+i) # GET URL OF THE NEEDED SECTION
            soup = bs4.BeautifulSoup(Url.text, 'html.parser') 
            span = soup.select("span#"+i) #SELECTION THE NECESSARY SPAN
            j = 0
            data = []
            text = ''
            text = span[0].find_all_next('p')[0].text # FINDING THE NEXT PARA FOR DATA
            span2 = span[0].find_all_next('p')[0]
            # UPDATING THE HEADINGS LIST
            headings += [span[0].text]
            for j in span2.find_next_siblings(): #FINDING CONSECUTIVE DATA
                if j.name == 'h2': # BREAK IF THE SECTION GETS OVER
                    break 
                if j.name == 'style':
                    continue
                text = text + " " + j.text #  ADDING TEXT TO THE PARA
                # REMOVING UNWANTED TEXT USING REGULAR EXPRESSIONS
                text = re.sub(r'\[.*\]','',text)
                text = re.sub(r'\\[^nt]','',text)
                text = re.sub(r'\\xa0','',text)
                text = re.sub(r':',':\n',text)
                text = re.sub(r'(\n)+','\n',text)
                data = text.split(' ')
                # KEEPING THE WORD LENGTH TO A MAX OF 800
                if(len(data)>int(800/len(AddSections))):
                    break 
            # UPDATING THE PARAGRAPHS LIST       
            paragraphs += [text]
    except:
        return [-1],[-1] # RETURN IN CASE OF ERROR
    return headings,paragraphs # RETURNS WHEN SUCCESSFUL

#FUNCTION TO MAKE THE REPORT OUT OF THE HEADINGS AND PARAGRAPHS 
def MakeReport(headings,paragraphs,fileName,heading,name,usn,secAlign,bodyAlign,color,underline,secSize,bodySize):  

    current = os.getcwd() # CURRENT WORKING DIRECTORY
    Base = os.path.join(current,'layout.docx') # PATH FOR THE SET LAYOUT DOC
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') # PATH FOR THE USERS DESKTOP
    if fileName.endswith(".docx"): #CHECK IF THE USER HAS ENTERED .DOCX IN THE FILE NAME
        output = os.path.join(desktop,fileName)
    else: # IF NOT APPEND .DOCX
        output = os.path.join(desktop,fileName+'.docx') 
    # CREATE A COPY OF THE LAYOUT FILE TO WRITE
    shutil.copyfile(Base,output)

    # OPEN THIS DOCUMENT FOR WRITING
    doc = docx.Document(output)

    # ADDING PARAGRAPH FOR THE REPORT TITLE
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(heading.upper())
    run.add_break()
    font = run.font
    font.size = Pt(30)
    if color ==  0:
        font.color.rgb = RGBColor(0xCC,0x0,0x0)
    elif color ==  1:
        font.color.rgb = RGBColor(0x0,0xCC,0x0)
    elif color ==  2:
        font.color.rgb = RGBColor(0x0,0x0,0xCC)
    elif color == 3:
        font.color.rgb = RGBColor(0x0,0x0,0x0)
    font.bold = True   
    
    # ADDING PARAGRAPH FOR THE 'SUBMITTED BY' CLAUSE
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Submitted by')
    run.add_break()
    font = run.font
    font.size = Pt(20)
    font.italic = True

    # ADDING PARAGRAPH FOR NAME AND USN 
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name+"\t("+usn+")")
    doc.add_page_break()
    font = run.font
    if color ==  0:
        font.color.rgb = RGBColor(0xCC,0x0,0x0)
    elif color ==  1:
        font.color.rgb = RGBColor(0x0,0xCC,0x0)
    elif color ==  2:
        font.color.rgb = RGBColor(0x0,0x0,0xCC)
    elif color == 3:
        font.color.rgb = RGBColor(0x0,0x0,0x0)
    font.size = Pt(20)
    font.bold = True  
    

    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Table Of Content')
    run.add_break()
    font = run.font
    font.underline = True
    font.size = Pt(30)
    if color ==  0:
        font.color.rgb = RGBColor(0xCC,0x0,0x0)
    elif color ==  1:
        font.color.rgb = RGBColor(0x0,0xCC,0x0)
    elif color ==  2:
        font.color.rgb = RGBColor(0x0,0x0,0xCC)
    elif color == 3:
        font.color.rgb = RGBColor(0x0,0x0,0x0)
    font.bold = True
    run.add_break()

    for i in range(0,len(headings)):
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(str(i+1)+"\t"+str(headings[i]))
        run.add_break()
        font = run.font
        font.size = Pt(20)
    doc.add_page_break()



    # MAIN CONTENT PART
    j = 0  

    for i in headings: # BROWSE THROUGH HEADINGS

        para = doc.add_paragraph() # ADDING PARAGRAPH FOR THESE HEADINGS WITH APPROPRIATE FORMATTING
        if secAlign == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif secAlign == 1:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif secAlign == 2:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif secAlign == 3:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = para.add_run(i)
        run.add_break()
        font = run.font
        font.size = Pt(secSize)
        if color ==  0:
            font.color.rgb = RGBColor(0xCC,0x0,0x0)
        elif color ==  1:
            font.color.rgb = RGBColor(0x0,0xCC,0x0)
        elif color ==  2:
            font.color.rgb = RGBColor(0x0,0x0,0xCC)
        elif color == 3:
            font.color.rgb = RGBColor(0x0,0x0,0x0)
        font.bold = True
        if underline == 0:
            font.underline = False
        elif underline == 1:
            font.underline = True



        para1 = doc.add_paragraph() # ADDING PARAGRAPH FOR THE CONTENT OF EACH HEADINGS WITH APPROPRIATE FORMATTING 

        if bodyAlign == 0:
            para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif bodyAlign == 1:
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif bodyAlign == 2:
            para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif bodyAlign == 3:
            para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run1 = para1.add_run(paragraphs[j])
        j += 1
        run1.add_break()
        font1 = run1.font
        font1.size = Pt(bodySize)

    doc.save(output) # SAVING THE DOCX FILE



   
        
    
